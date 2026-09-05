"""
Insider Trading Daily Report Generator
=======================================
Reads insider_trades_1m_plus.csv and generates a daily report with:
  1. All insider trades per day for the last 5 days
  2. In-depth purchase cluster analysis (entire history)
  3. Additional insights: sentiment, top buyers, repeat buyer alerts

Output: Markdown (.md) + Word document (.docx)

Usage:
  python report_gen.py
"""

import pandas as pd
import os
import sys
import re
import markdown
import yfinance as yf
import matplotlib
matplotlib.use('Agg')  # non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime, timedelta
from pathlib import Path

# Word doc imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

# Fix Windows console encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')


# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION — Easy to modify
# ═════════════════════════════════════════════════════════════════════════════

CSV_FILE = "insider_trades_1m_plus.csv"
LOOKBACK_DAYS = 10
CHART_DIR = "cluster_charts"  # subfolder for performance charts

# --- Section 2: Purchase Cluster Analysis filters ---
# Keywords to identify purchase transactions (matched case-insensitively in Text)
PURCHASE_KEYWORDS = ["purchase"]  # Add more: e.g. ["purchase", "buy", "acquisition"]

# Minimum number of distinct insiders who must have bought on the same day/ticker
MIN_INSIDERS = 2

# Minimum dollar value per individual insider trade
MIN_VALUE_PER_INSIDER = 1_000_000  # $3M


# ═════════════════════════════════════════════════════════════════════════════
# DATA LOADING & PREPARATION
# ═════════════════════════════════════════════════════════════════════════════

def load_and_prepare(csv_path: str) -> pd.DataFrame:
    """Load CSV, clean types, classify trades from Text column."""
    df = pd.read_csv(csv_path)

    # Parse Value: strip commas → float
    df["Value_num"] = (
        df["Value"]
        .astype(str)
        .str.replace(",", "", regex=False)
        .apply(pd.to_numeric, errors="coerce")
    )

    # Parse dates
    df["Date"] = pd.to_datetime(df["Start Date"], errors="coerce")

    # Re-classify Trade_Type from Text (the original classifier used the
    # empty Transaction column, so everything was "Other")
    df["Trade_Type"] = df["Text"].apply(_classify_from_text)

    # Absolute value for sorting
    df["Abs_Value"] = df["Value_num"].abs()

    return df


def _classify_from_text(text: str) -> str:
    """Classify a trade as Buy/Sell/Other based on the Text column."""
    t = str(text).lower()
    if any(kw in t for kw in ("purchase", "buy", "acquisition")):
        return "Buy"
    elif any(kw in t for kw in ("sale", "sell", "disposition")):
        return "Sell"
    elif any(kw in t for kw in ("conversion", "exercise")):
        return "Exercise/Conversion"
    elif any(kw in t for kw in ("grant", "award")):
        return "Grant/Award"
    elif "gift" in t:
        return "Gift"
    return "Other"


def _sentiment_label(trade_type: str) -> str:
    """Map trade type to a sentiment emoji+label."""
    mapping = {
        "Buy": "🟢 Bullish",
        "Sell": "🔴 Bearish",
        "Exercise/Conversion": "🟡 Neutral",
        "Grant/Award": "🟡 Neutral",
        "Gift": "🟡 Neutral",
        "Other": "⚪ Unknown",
    }
    return mapping.get(trade_type, "⚪ Unknown")


def _fmt_value(val) -> str:
    """Format a numeric value as $X,XXX,XXX."""
    try:
        return f"${val:,.0f}"
    except (ValueError, TypeError):
        return str(val)


def _is_purchase(text: str) -> bool:
    """Check if a text matches any PURCHASE_KEYWORDS."""
    t = str(text).lower()
    return any(kw in t for kw in PURCHASE_KEYWORDS)


# ═════════════════════════════════════════════════════════════════════════════
# REPORT SECTION BUILDERS
# ═════════════════════════════════════════════════════════════════════════════

def build_section0_all_buys(df: pd.DataFrame, today: datetime, chart_dir: Path) -> str:
    """
    Top section: All buy/purchase transactions in the last N days,
    sorted by value descending in a single table.
    Includes a return chart for each unique ticker.
    """
    cutoff = today - timedelta(days=LOOKBACK_DAYS)
    recent = df[(df["Date"] >= cutoff) & (df["Trade_Type"] == "Buy")].copy()

    md = f"## \U0001f7e2 All Insider Purchases (Last {LOOKBACK_DAYS} Days)\n\n"

    if recent.empty:
        md += f"_No purchase transactions found in the last {LOOKBACK_DAYS} days._\n\n"
        return md

    recent = recent.sort_values("Abs_Value", ascending=False)
    total_val = recent["Abs_Value"].sum()
    n_trades = len(recent)
    n_tickers = recent["Ticker"].nunique()

    md += (
        f"**{n_trades} purchase(s)** across **{n_tickers} ticker(s)** "
        f"totalling **{_fmt_value(total_val)}**\n\n"
    )
    md += _df_to_md_table(
        recent,
        columns=["Ticker", "Insider", "Position", "Shares", "Value", "Start Date", "Text"],
        headers=["Ticker", "Insider", "Position", "Shares", "Value ($)", "Date", "Details"],
    )
    md += "\n"

    # Generate a 1-year return chart for each unique ticker
    one_year_ago = (today - timedelta(days=365)).date()
    unique_tickers = recent["Ticker"].unique()
    for ticker in unique_tickers:
        print(f"       Chart: {ticker} (1-year return)...", end=" ", flush=True)
        chart_path = _plot_cluster_return(ticker, one_year_ago, chart_dir)
        if chart_path:
            md += f"![{ticker} vs S&P 500 — 1-Year Return]({chart_path})\n\n"
            print("\u2705")
        else:
            md += f"_Chart unavailable for {ticker}._\n\n"
            print("\u2014")

    return md


def build_section1_daily_trades(df: pd.DataFrame, today: datetime) -> str:
    """
    Section 1: All insider trades per day for the last N days.
    Separate table per day, ordered by biggest trade to smallest.
    Days ordered most recent first.
    """
    cutoff = today - timedelta(days=LOOKBACK_DAYS)
    recent = df[df["Date"] >= cutoff].copy()

    if recent.empty:
        return (
            "## 📊 Section 1: Daily Insider Trades (Last 5 Days)\n\n"
            "_No trades found in the last 5 days._\n\n"
        )

    md = "## 📊 Section 1: Daily Insider Trades (Last 5 Days)\n\n"
    md += (
        f"Showing all insider trades ≥ $1M from "
        f"**{cutoff.strftime('%Y-%m-%d')}** to **{today.strftime('%Y-%m-%d')}**, "
        f"ordered by transaction size within each day.\n\n"
    )

    # Group by date, sorted descending
    recent["DateOnly"] = recent["Date"].dt.date
    dates_sorted = sorted(recent["DateOnly"].dropna().unique(), reverse=True)

    for d in dates_sorted:
        day_df = recent[recent["DateOnly"] == d].sort_values("Abs_Value", ascending=False)
        total_val = day_df["Abs_Value"].sum()
        n_trades = len(day_df)

        md += f"### 📅 {d} — {n_trades} trade(s), {_fmt_value(total_val)} total\n\n"
        md += _df_to_md_table(
            day_df,
            columns=["Ticker", "Insider", "Position", "Trade_Type", "Shares", "Value", "Text"],
            headers=["Ticker", "Insider", "Position", "Type", "Shares", "Value ($)", "Details"],
        )
        md += "\n"

    return md


def _plot_cluster_return(ticker: str, start_date, chart_dir: Path) -> str | None:
    """
    Plot cumulative % return for `ticker` vs S&P 500 from `start_date` to today.
    Saves the chart as a PNG and returns the file path, or None on failure.
    """
    try:
        start = pd.Timestamp(start_date)
        end = pd.Timestamp(datetime.now())

        # Need at least a few days of data
        if (end - start).days < 2:
            return None

        # Download price data
        tk_data = yf.download(ticker, start=start, end=end, progress=False, auto_adjust=True)
        sp_data = yf.download("^GSPC", start=start, end=end, progress=False, auto_adjust=True)

        if tk_data.empty or sp_data.empty or len(tk_data) < 2:
            return None

        # Handle MultiIndex columns from yfinance
        if isinstance(tk_data.columns, pd.MultiIndex):
            tk_close = tk_data["Close"].iloc[:, 0]
        else:
            tk_close = tk_data["Close"]

        if isinstance(sp_data.columns, pd.MultiIndex):
            sp_close = sp_data["Close"].iloc[:, 0]
        else:
            sp_close = sp_data["Close"]

        # Calculate cumulative % return
        tk_return = (tk_close / tk_close.iloc[0] - 1) * 100
        sp_return = (sp_close / sp_close.iloc[0] - 1) * 100

        # Final return values for annotation
        tk_final = tk_return.iloc[-1]
        sp_final = sp_return.iloc[-1]

        # Plot
        fig, ax = plt.subplots(figsize=(10, 5))
        fig.patch.set_facecolor('#1a1a2e')
        ax.set_facecolor('#16213e')

        ax.plot(tk_return.index, tk_return.values, color='#00d2ff', linewidth=2,
                label=f'{ticker} ({tk_final:+.1f}%)')
        ax.plot(sp_return.index, sp_return.values, color='#ff6b6b', linewidth=2,
                alpha=0.8, label=f'S&P 500 ({sp_final:+.1f}%)')

        ax.axhline(y=0, color='#555555', linewidth=0.8, linestyle='--')

        ax.set_title(
            f'{ticker} vs S&P 500 — Return Since Insider Purchase ({start.strftime("%Y-%m-%d")})',
            color='white', fontsize=13, fontweight='bold', pad=12
        )
        ax.set_ylabel('Cumulative Return (%)', color='white', fontsize=11)
        ax.set_xlabel('', color='white')

        ax.tick_params(colors='white', labelsize=9)
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %Y'))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator())
        fig.autofmt_xdate(rotation=30)

        ax.legend(loc='best', fontsize=10, facecolor='#16213e', edgecolor='#444',
                  labelcolor='white')
        ax.grid(True, alpha=0.15, color='white')

        for spine in ax.spines.values():
            spine.set_color('#333')

        plt.tight_layout()

        # Save
        chart_dir.mkdir(exist_ok=True)
        safe_date = str(start_date).replace(' ', '_').replace(':', '-')
        filename = f"{ticker}_{safe_date}.png"
        filepath = chart_dir / filename
        fig.savefig(filepath, dpi=140, bbox_inches='tight', facecolor=fig.get_facecolor())
        plt.close(fig)

        return str(filepath)

    except Exception as e:
        print(f"       ⚠️  Chart failed for {ticker}: {e}")
        return None


def build_section2_purchase_clusters(df: pd.DataFrame, chart_dir: Path) -> str:
    """
    Section 2: In-depth purchase cluster analysis (entire history).
    Find cases where >MIN_INSIDERS insiders bought >MIN_VALUE_PER_INSIDER each
    on the same day for the same ticker.
    Includes a performance chart (ticker vs S&P 500) under each cluster.
    """
    # Filter to purchases
    purchases = df[df["Text"].apply(_is_purchase)].copy()

    if purchases.empty:
        return (
            "## 🔍 Section 2: In-Depth Purchase Cluster Analysis\n\n"
            "_No purchase transactions found._\n\n"
        )

    # Filter by minimum value per insider
    big_purchases = purchases[purchases["Abs_Value"] >= MIN_VALUE_PER_INSIDER].copy()
    big_purchases["DateOnly"] = big_purchases["Date"].dt.date

    if big_purchases.empty:
        return (
            "## 🔍 Section 2: In-Depth Purchase Cluster Analysis\n\n"
            f"_No purchases ≥ {_fmt_value(MIN_VALUE_PER_INSIDER)} found._\n\n"
        )

    # Group by (date, ticker) and count distinct insiders
    groups = big_purchases.groupby(["DateOnly", "Ticker"])

    clusters = []
    for (date, ticker), group in groups:
        unique_insiders = group["Insider"].nunique()
        if unique_insiders >= MIN_INSIDERS:
            clusters.append({
                "date": date,
                "ticker": ticker,
                "n_insiders": unique_insiders,
                "total_value": group["Abs_Value"].sum(),
                "data": group.sort_values("Abs_Value", ascending=False),
            })

    if not clusters:
        return (
            "## 🔍 Section 2: In-Depth Purchase Cluster Analysis\n\n"
            f"_No cases found where ≥{MIN_INSIDERS} insiders each bought "
            f"≥{_fmt_value(MIN_VALUE_PER_INSIDER)} on the same day._\n\n"
        )

    # Sort: date descending, then total value descending
    clusters.sort(key=lambda c: (-c["date"].toordinal(), -c["total_value"]))

    min_val_m = MIN_VALUE_PER_INSIDER / 1_000_000
    md = "## 🔍 Section 2: In-Depth Purchase Cluster Analysis\n\n"
    md += (
        f"Scanning **entire history** for cases where ≥ **{MIN_INSIDERS} insiders** "
        f"each purchased ≥ **${min_val_m:.0f}M** of the same stock on the same day.\n\n"
        f"**{len(clusters)} cluster(s) found.**\n\n"
    )

    for i, c in enumerate(clusters, 1):
        total_m = c["total_value"] / 1_000_000
        md += (
            f"### 📅 {c['date']} — {c['ticker']} "
            f"({c['n_insiders']} insiders, ${total_m:,.1f}M total)\n\n"
        )
        md += _df_to_md_table(
            c["data"],
            columns=["Ticker", "Insider", "Position", "Shares", "Value", "Ownership", "Text"],
            headers=["Ticker", "Insider", "Position", "Shares", "Value ($)", "Ownership", "Details"],
        )
        md += "\n"

        # Generate performance chart
        print(f"       Chart [{i}/{len(clusters)}]: {c['ticker']} from {c['date']}...",
              end=" ", flush=True)
        chart_path = _plot_cluster_return(c["ticker"], c["date"], chart_dir)
        if chart_path:
            md += f"![{c['ticker']} vs S&P 500 since {c['date']}]({chart_path})\n\n"
            print("✅")
        else:
            md += f"_Chart unavailable for {c['ticker']} (insufficient data)._\n\n"
            print("—")

    return md


def build_section3_insights(df: pd.DataFrame, today: datetime) -> str:
    """
    Section 3: Additional insights — sentiment breakdown, top buyers,
    repeat buyer alerts.
    """
    cutoff = today - timedelta(days=LOOKBACK_DAYS)
    recent = df[df["Date"] >= cutoff].copy()

    md = "## 📈 Section 3: Additional Insights (Last 5 Days)\n\n"

    if recent.empty:
        md += "_No data for the last 5 days._\n\n"
        return md

    # ── 3a: Trade Type Breakdown ──────────────────────────────────────────
    md += "### Trade Type Breakdown\n\n"
    type_summary = (
        recent.groupby("Trade_Type")
        .agg(
            Count=("Trade_Type", "size"),
            Total_Value=("Abs_Value", "sum"),
        )
        .sort_values("Total_Value", ascending=False)
        .reset_index()
    )
    type_summary["Avg_Value"] = type_summary["Total_Value"] / type_summary["Count"]
    type_summary["Total_Value"] = type_summary["Total_Value"].apply(_fmt_value)
    type_summary["Avg_Value"] = type_summary["Avg_Value"].apply(_fmt_value)
    type_summary.columns = ["Trade Type", "# Trades", "Total Value", "Avg Trade"]

    md += _raw_df_to_md_table(type_summary)
    md += "\n"

    # ── 3b: Sentiment Summary by Ticker ───────────────────────────────────
    md += "### Sentiment Summary by Ticker\n\n"
    recent["Sentiment"] = recent["Trade_Type"].apply(_sentiment_label)

    sent_summary = (
        recent.groupby(["Ticker", "Sentiment"])
        .agg(Count=("Sentiment", "size"), Volume=("Abs_Value", "sum"))
        .reset_index()
        .sort_values(["Ticker", "Volume"], ascending=[True, False])
    )
    sent_summary["Volume"] = sent_summary["Volume"].apply(_fmt_value)
    sent_summary.columns = ["Ticker", "Sentiment", "# Trades", "Volume"]

    # Only show tickers with some buy/sell activity (skip if too many)
    tickers_with_buys = recent[recent["Trade_Type"] == "Buy"]["Ticker"].unique()
    if len(tickers_with_buys) > 0:
        highlight = sent_summary[sent_summary["Ticker"].isin(tickers_with_buys)]
        md += f"_Showing tickers with at least one purchase ({len(tickers_with_buys)} tickers):_\n\n"
        md += _raw_df_to_md_table(highlight)
    else:
        md += "_No purchase activity in the last 5 days._\n\n"
        # Show top 10 tickers by volume instead
        top_tickers = (
            recent.groupby("Ticker")["Abs_Value"].sum()
            .nlargest(10).index
        )
        top_sent = sent_summary[sent_summary["Ticker"].isin(top_tickers)]
        md += f"_Top 10 tickers by volume:_\n\n"
        md += _raw_df_to_md_table(top_sent)
    md += "\n"

    # ── 3c: Top Buyers (by purchase volume) ───────────────────────────────
    md += "### 🏆 Top Buyers (Last 5 Days)\n\n"
    buyers = recent[recent["Trade_Type"] == "Buy"].copy()
    if not buyers.empty:
        top_buyers = (
            buyers.groupby(["Insider", "Ticker"])
            .agg(
                Trades=("Insider", "size"),
                Total_Value=("Abs_Value", "sum"),
                Position=("Position", "first"),
            )
            .sort_values("Total_Value", ascending=False)
            .reset_index()
            .head(15)
        )
        top_buyers["Total_Value"] = top_buyers["Total_Value"].apply(_fmt_value)
        top_buyers.columns = ["Insider", "Ticker", "# Trades", "Total Value", "Position"]
        top_buyers = top_buyers[["Insider", "Ticker", "Position", "# Trades", "Total Value"]]
        md += _raw_df_to_md_table(top_buyers)
    else:
        md += "_No purchases in the last 5 days._\n"
    md += "\n"

    # ── 3d: Repeat Buyer Alert ────────────────────────────────────────────
    md += "### 🔔 Repeat Buyer Alert\n\n"
    md += "_Insiders who purchased the same stock on multiple days in the last 5 days:_\n\n"
    if not buyers.empty:
        buyers["DateOnly"] = buyers["Date"].dt.date
        repeat = (
            buyers.groupby(["Insider", "Ticker"])
            .agg(
                Days=("DateOnly", "nunique"),
                Total_Trades=("Insider", "size"),
                Total_Value=("Abs_Value", "sum"),
                Dates=("DateOnly", lambda x: ", ".join(str(d) for d in sorted(x.unique()))),
            )
            .reset_index()
        )
        repeat = repeat[repeat["Days"] > 1].sort_values("Total_Value", ascending=False)
        if not repeat.empty:
            repeat["Total_Value"] = repeat["Total_Value"].apply(_fmt_value)
            repeat.columns = ["Insider", "Ticker", "# Days", "# Trades", "Total Value", "Dates"]
            md += _raw_df_to_md_table(repeat)
        else:
            md += "_No repeat buyers found._\n"
    else:
        md += "_No purchase data available._\n"
    md += "\n"

    # ── 3e: Largest Individual Trades ─────────────────────────────────────
    md += "### 💰 Top 10 Largest Individual Trades (Last 5 Days)\n\n"
    top10 = recent.nlargest(10, "Abs_Value")
    md += _df_to_md_table(
        top10,
        columns=["Ticker", "Insider", "Position", "Trade_Type", "Shares", "Value", "Start Date", "Text"],
        headers=["Ticker", "Insider", "Position", "Type", "Shares", "Value ($)", "Date", "Details"],
    )
    md += "\n"

    return md


# ═════════════════════════════════════════════════════════════════════════════
# MARKDOWN TABLE HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def _df_to_md_table(df: pd.DataFrame, columns: list, headers: list) -> str:
    """Convert selected DataFrame columns to a markdown table."""
    # Build header
    md = "| " + " | ".join(headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

    # Build rows
    for _, row in df.iterrows():
        cells = []
        for col in columns:
            val = row.get(col, "")
            # Clean up for markdown (escape pipes)
            val = str(val).replace("|", "\\|").replace("\n", " ")
            cells.append(val)
        md += "| " + " | ".join(cells) + " |\n"

    return md


def _raw_df_to_md_table(df: pd.DataFrame) -> str:
    """Convert an entire DataFrame to a markdown table."""
    cols = list(df.columns)
    md = "| " + " | ".join(cols) + " |\n"
    md += "| " + " | ".join(["---"] * len(cols)) + " |\n"
    for _, row in df.iterrows():
        cells = [str(row[c]).replace("|", "\\|").replace("\n", " ") for c in cols]
        md += "| " + " | ".join(cells) + " |\n"
    return md


# ═════════════════════════════════════════════════════════════════════════════
# MARKDOWN → WORD CONVERSION
# ═════════════════════════════════════════════════════════════════════════════

def markdown_to_word(md_content: str, output_file: str):
    """Convert markdown string to a formatted Word document."""
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            # End table if we were in one
            if in_table:
                in_table = False
                _render_word_table(doc, table_rows)
                table_rows = []
            i += 1
            continue

        # Table rows
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Skip separator rows (---  rows)
            if not all(set(c.replace('-', '').replace(':', '').strip()) == set() for c in cells):
                table_rows.append(cells)

            i += 1
            continue

        # End of table
        if in_table and '|' not in stripped:
            in_table = False
            _render_word_table(doc, table_rows)
            table_rows = []

        # Headers
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            heading = doc.add_heading(text, level=min(level, 9))
            if level == 1:
                heading.style.font.color.rgb = RGBColor(0, 51, 102)
                heading.style.font.size = Pt(22)
            elif level == 2:
                heading.style.font.color.rgb = RGBColor(0, 102, 204)
                heading.style.font.size = Pt(16)

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            doc.add_paragraph(text, style='List Bullet')

        # Images: ![alt text](path)
        elif stripped.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                if os.path.exists(image_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(image_path, width=Inches(6))
                    if alt_text:
                        cap = doc.add_paragraph(alt_text)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cap.runs:
                            r.font.size = Pt(9)
                            r.font.italic = True
                            r.font.color.rgb = RGBColor(100, 100, 100)
                else:
                    p = doc.add_paragraph(f"[Image not found: {alt_text}]")
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Italic block (lines starting with _..._)
        elif stripped.startswith('_') and stripped.endswith('_'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('_'))
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Bold/italic/regular paragraphs
        else:
            p = doc.add_paragraph()
            parts = stripped.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 0:
                    italic_parts = part.split('*')
                    for i_idx, i_part in enumerate(italic_parts):
                        run = p.add_run(i_part)
                        if i_idx % 2 == 1:
                            run.italic = True
                else:
                    run = p.add_run(part)
                    run.bold = True

        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _render_word_table(doc, table_rows)

    doc.save(output_file)
    print(f"  ✅ Word document: {output_file}")


def _render_word_table(doc, table_rows):
    """Render a markdown-style table into a Word document."""
    if not table_rows:
        return

    num_rows = len(table_rows)
    num_cols = len(table_rows[0])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for row_idx, row_data in enumerate(table_rows):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = cell_data

                # Header row styling
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(
                            r'<w:shd {} w:fill="0066CC"/>'.format(
                                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                            )
                        )
                    )

                # Font size for all cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    doc.add_paragraph()  # spacing after table


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

def main():
    print()
    print("=" * 65)
    print("  📋 Insider Trading Daily Report Generator")
    print("=" * 65)

    # Resolve CSV path relative to this script
    script_dir = Path(__file__).resolve().parent
    csv_path = script_dir / CSV_FILE

    if not csv_path.exists():
        print(f"  ❌ CSV not found: {csv_path}")
        print(f"     Run sp500_insider_trades.py first to generate the data.")
        sys.exit(1)

    # Load data
    print(f"\n  📂 Loading {CSV_FILE}...")
    df = load_and_prepare(str(csv_path))
    print(f"     {len(df)} trades loaded ({df['Ticker'].nunique()} tickers)")
    print(f"     Date range: {df['Date'].min():%Y-%m-%d} → {df['Date'].max():%Y-%m-%d}")

    today = datetime.now()
    date_str = today.strftime("%Y-%m-%d")

    # Prepare chart directory
    chart_dir = script_dir / CHART_DIR
    chart_dir.mkdir(exist_ok=True)

    # Build report sections
    print(f"\n  📝 Building report (reference date: {date_str})...")

    # Title
    report = f"# 📊 Insider Trading Report — {date_str}\n\n"
    report += (
        f"**Generated:** {today.strftime('%Y-%m-%d %H:%M')}  \n"
        f"**Data source:** {CSV_FILE}  \n"
        f"**Total trades in database:** {len(df):,}  \n"
        f"**Date range:** {df['Date'].min():%Y-%m-%d} → {df['Date'].max():%Y-%m-%d}  \n"
        f"**Lookback window:** {LOOKBACK_DAYS} days  \n\n"
    )
    report += "---\n\n"

    # Top section: All buys in the lookback window
    print("     All Buys summary + charts...")
    report += build_section0_all_buys(df, today, chart_dir)
    report += "---\n\n"

    # Section 1: Daily trades (last N days)
    print("     Section 1: Daily insider trades...")
    report += build_section1_daily_trades(df, today)
    report += "---\n\n"

    # Section 2: Purchase cluster analysis (entire history) + charts
    print("     Section 2: Purchase cluster analysis + charts...")
    report += build_section2_purchase_clusters(df, chart_dir)
    report += "---\n\n"

    # Section 3: Additional insights (last 5 days)
    print("     Section 3: Additional insights...")
    report += build_section3_insights(df, today)

    # Footer
    report += "---\n\n"
    report += (
        f"_Report generated on {today.strftime('%Y-%m-%d %H:%M')} by report_gen.py. "
        f"Data filtered to trades ≥ $1M._\n"
    )

    # Save markdown
    md_file = script_dir / f"insider_report_{date_str}.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(report)
    print(f"\n  ✅ Markdown report: {md_file.name}")

    # Convert to Word
    docx_file = script_dir / f"insider_report_{date_str}.docx"
    print(f"  📄 Converting to Word...")
    try:
        markdown_to_word(report, str(docx_file))
    except PermissionError:
        # File is likely open in Word — use a timestamped name
        ts = today.strftime("%H%M%S")
        docx_file = script_dir / f"insider_report_{date_str}_{ts}.docx"
        print(f"  ⚠️  Original file locked, saving as: {docx_file.name}")
        markdown_to_word(report, str(docx_file))

    print()
    print("=" * 65)
    print(f"  ✅ DONE! Files saved in: {script_dir}")
    print(f"     📄 {md_file.name}")
    print(f"     📄 {docx_file.name}")
    print("=" * 65)
    print()


if __name__ == "__main__":
    main()
